from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "alter_domus_fund_reporting_control_tower_pitch.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
MUTED = RGBColor(85, 95, 110)
BLACK = RGBColor(0, 0, 0)
LIGHT_GRAY = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
BORDER = "D9DEE7"


def dxa(inches):
    return int(round(inches * 1440))


def set_run_font(run, size=11, color=BLACK, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size, color=BLACK, bold=False, name="Calibri"):
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold


def configure_document(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_style_font(normal, 11, BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, 16, BLUE, bold=True)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, 13, BLUE, bold=True)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    set_style_font(h3, 12, DARK_BLUE, bold=True)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    for list_style_name in ("List Bullet", "List Number"):
        style = doc.styles[list_style_name]
        set_style_font(style, 11, BLACK)
        pf = style.paragraph_format
        pf.left_indent = Inches(0.5)
        pf.first_line_indent = Inches(-0.25)
        pf.space_after = Pt(8)
        pf.line_spacing = 1.167

    header_p = section.header.paragraphs[0]
    header_p.text = ""
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_p.add_run("Internal Idea Submission | Fund Reporting Control Tower")
    set_run_font(header_run, size=9, color=MUTED)

    footer_p = section.footer.paragraphs[0]
    footer_p.text = ""
    footer_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer_run = footer_p.add_run("Prepared for Alter Domus innovation review")
    set_run_font(footer_run, size=9, color=MUTED)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        elem = tc_mar.find(qn(f"w:{side}"))
        if elem is None:
            elem = OxmlElement(f"w:{side}")
            tc_mar.append(elem)
        elem.set(qn("w:w"), str(value))
        elem.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        elem = borders.find(tag)
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_fixed_table(table, widths, indent=120, header_fill=LIGHT_GRAY):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(1, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    set_table_borders(table)
    for row_index, row in enumerate(table.rows):
        if row_index == 0 and header_fill:
            mark_header_row(row)
        for index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            if row_index == 0 and header_fill:
                shade_cell(cell, header_fill)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.10


def clear_cell(cell):
    cell.text = ""
    return cell.paragraphs[0]


def write_cell(cell, text, bold=False, color=BLACK, size=10.5):
    p = clear_cell(cell)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    return p


def add_title_block(doc):
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(3)
    run = kicker.add_run("IDEA SUBMISSION")
    set_run_font(run, size=10, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.line_spacing = 1.05
    run = title.add_run("Fund Reporting Control Tower")
    set_run_font(run, size=24, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run(
        "An AI-assisted, audit-ready workbench for faster and safer fund administration reporting"
    )
    set_run_font(run, size=13, color=MUTED)

    meta = doc.add_table(rows=4, cols=2)
    set_fixed_table(meta, [dxa(1.55), dxa(4.95)], header_fill=None)
    rows = [
        ("Audience", "Alter Domus internal innovation and resource allocation review"),
        ("Decision requested", "Approve a 6-8 week pilot using sanitized fund data and a small cross-functional team"),
        ("Prototype basis", "FA Reporting Studio application: fund repository, template mapping, cash-flow generation, validation, lineage, and controlled agent tooling"),
        ("Prepared", "July 5, 2026"),
    ]
    for row, (label, value) in zip(meta.rows, rows):
        shade_cell(row.cells[0], LIGHT_GRAY)
        write_cell(row.cells[0], label, bold=True, color=DARK_BLUE, size=10.2)
        write_cell(row.cells[1], value, size=10.2)


def add_callout(doc, heading, body):
    table = doc.add_table(rows=1, cols=1)
    set_fixed_table(table, [9360], header_fill=None)
    cell = table.cell(0, 0)
    shade_cell(cell, CALLOUT_FILL)
    p = clear_cell(cell)
    p.paragraph_format.space_after = Pt(3)
    head = p.add_run(heading)
    set_run_font(head, size=11.5, color=NAVY, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    run = p2.add_run(body)
    set_run_font(run, size=10.5, color=BLACK)


def add_paragraph(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size={1: 16, 2: 13, 3: 12}.get(level, 11), color=BLUE if level < 3 else DARK_BLUE, bold=True)
    return p


def add_three_col_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=3)
    set_fixed_table(table, widths, header_fill=LIGHT_GRAY)
    for i, header in enumerate(headers):
        write_cell(table.rows[0].cells[i], header, bold=True, color=DARK_BLUE, size=9.8)
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            write_cell(cells[i], value, size=9.6)
            set_cell_width(cells[i], widths[i])
            set_cell_margins(cells[i])
    set_fixed_table(table, widths, header_fill=LIGHT_GRAY)
    return table


def build_doc():
    doc = Document()
    configure_document(doc)
    add_title_block(doc)

    add_callout(
        doc,
        "One-sentence pitch",
        "Build an internal control layer that turns fund source documents, client workbooks, TB/GL files, and human approvals into repeatable, traceable, validation-ready reporting workflows.",
    )

    add_heading(doc, "The Pitch", 1)
    add_paragraph(
        doc,
        "Alter Domus already differentiates through alternative-fund expertise, global operating scale, and technology-enabled service delivery. The next opportunity is to make that expertise more repeatable inside the reporting process: every fund team should be able to reuse governed source knowledge, approved mappings, validation rules, and audit evidence instead of recreating the same judgment from scattered files each period.",
    )
    add_paragraph(
        doc,
        "The app we are building is a prototype of that operating layer. It does not try to remove fund accountants from the process. It gives them a controlled workspace where automation prepares the draft, explains the source support, flags uncertainty, and blocks final export until review and validation pass.",
    )

    add_heading(doc, "Problem Worth Solving", 1)
    for item in [
        "Client reporting is template-heavy: each manager, fund, vehicle, and period can have different workbook expectations.",
        "Key terms live across LPAs, amendments, side letters, capital notices, NAV packages, bank files, audit support, and manual notes.",
        "Mapping TB/GL activity to report lines is repetitive but risky; a small mapping error can become a client issue, audit question, or rework loop.",
        "Evidence is often separated from the output: reviewers need to know which source file, version, account, or approved mapping supports each number.",
        "AI is useful only if it is bounded by fund-admin controls: human review, source grounding, export approval, audit trail, and data governance.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Product Concept", 1)
    add_three_col_table(
        doc,
        ["Layer", "What the app does", "Why it matters"],
        [
            (
                "Fund Repository",
                "Stores versioned fund documents and datasets, reads current sources, extracts key points, and lets reviewers confirm, correct, or dismiss facts.",
                "Creates a reusable fund knowledge base with evidence, not just a file dump.",
            ),
            (
                "Semantic Mapping",
                "Maps source accounts and client template rows into a canonical reporting vocabulary with deterministic suggestions and review status.",
                "Reduces one-off workbook logic and makes approved reporting meaning portable across periods.",
            ),
            (
                "Report Builder",
                "Runs cash-flow workbook generation from stored or uploaded trial balance, general ledger, and active template inputs.",
                "Turns repeat reporting into a controlled workflow while preserving client-format workbooks.",
            ),
            (
                "Validation and Lineage",
                "Checks approved mappings, source grounding, unresolved rows, formula support, critical cash concepts, and auditability; persists line-level evidence.",
                "Supports reviewer confidence, audit support, and faster issue resolution.",
            ),
            (
                "Agent-Safe Workflow",
                "Provides a restricted tool catalog for assistants: draft work, analyze sources, suggest mappings, run validations, and request final export approval.",
                "Lets teams explore AI assistance without allowing uncontrolled finalization or hidden changes.",
            ),
        ],
        [dxa(1.35), dxa(2.75), dxa(2.4)],
    )

    add_heading(doc, "What the Prototype Already Demonstrates", 1)
    for item in [
        "A React admin Reporting Studio with sections for Run Report, Templates & Mapping, Review Issues, Report History, and Fund Repository.",
        "Specialist fund-source readers for LPAs, side letters, subscription agreements, capital notices, bank statements, NAV packages, financial statements, audit reports, valuation files, waterfall statements, and more.",
        "Versioned repository items with current-source analysis, extracted knowledge, review queues, source opening, and confirmed-fact conflict detection.",
        "Cash-flow template analysis with activation readiness, semantic coverage, anchor review, and missing-row preflight before workbook generation.",
        "Report runs that store input artifacts, attach confirmed repository knowledge, generate workbook outputs, and expose warnings, low-confidence mappings, and history.",
        "Validation services and export controls that require a ready validation result and an approval review task before final export.",
        "Lineage services that preserve source file identity, repository version hashes, mapping snapshots, confidence, and evidence for report rows and extractor assignments.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Why This Fits Alter Domus", 1)
    add_paragraph(
        doc,
        "This is a fund-administration productivity idea, but the strategic value is broader: it can strengthen Alter Domus's ability to scale specialized judgment across clients, jurisdictions, and reporting teams while preserving the review discipline that clients and auditors expect.",
    )
    add_three_col_table(
        doc,
        ["Strategic fit", "Potential impact", "Control posture"],
        [
            (
                "Alternative fund specialization",
                "Encodes fund-specific reporting terms, investor activity, capital activity, valuation context, and document evidence.",
                "Human reviewers confirm extracted terms before they become trusted knowledge.",
            ),
            (
                "Technology-enabled operations",
                "Moves recurring reporting work from manual assembly toward repeatable workflows with reusable mappings and templates.",
                "Automation produces drafts and exceptions; people approve final outputs.",
            ),
            (
                "Audit-ready service delivery",
                "Preserves source versions, report inputs, mapping decisions, validation checks, and export approvals.",
                "Final export is blocked unless readiness checks pass and an approval task is completed.",
            ),
            (
                "Client experience",
                "Shortens rework loops and enables clearer explanations when a client asks where a number came from.",
                "Each output can be backed by lineage rather than after-the-fact reconstruction.",
            ),
        ],
        [dxa(1.65), dxa(2.6), dxa(2.25)],
    )

    doc.add_page_break()
    add_heading(doc, "Pilot Proposal", 1)
    add_callout(
        doc,
        "Resource ask",
        "Approve a 6-8 week pilot with sanitized source files from 2-3 representative funds, one fund accounting SME, one reporting/review SME, one product owner, one engineer, and security/privacy input before any real client data is used.",
    )
    add_three_col_table(
        doc,
        ["Workstream", "Pilot output", "Decision gate"],
        [
            (
                "Fund source repository",
                "Load sanitized LPAs, side letters, NAV packages, TB/GL files, and reporting workbooks; extract and review key facts.",
                "Can reviewers trust and correct extracted facts quickly?",
            ),
            (
                "Template and mapping workflow",
                "Analyze 2-3 real reporting templates and approve semantic mappings for recurring cash-flow/report lines.",
                "Does mapping reuse reduce setup time without losing control?",
            ),
            (
                "Report generation",
                "Generate draft cash-flow workbooks from TB/GL inputs and preserve source evidence and warnings.",
                "Can the output support a real preparer-reviewer process?",
            ),
            (
                "Validation and export control",
                "Run readiness checks, create review tasks, and require approval before final export.",
                "Are the controls strong enough for internal governance review?",
            ),
        ],
        [dxa(1.7), dxa(2.75), dxa(2.05)],
    )

    add_heading(doc, "Pilot Success Measures", 1)
    add_three_col_table(
        doc,
        ["Measure", "Pilot target", "Why it matters"],
        [
            (
                "Report assembly effort",
                "Demonstrate a measurable reduction in preparer time for repeat cash-flow reporting after mappings are approved.",
                "Direct operating leverage.",
            ),
            (
                "Mapping coverage",
                "Reach high first-pass coverage on selected templates, with all low-confidence items routed to review.",
                "Productivity without hiding uncertainty.",
            ),
            (
                "Lineage completeness",
                "Every generated report line has source evidence or an explicit unresolved exception.",
                "Auditability and reviewer confidence.",
            ),
            (
                "Exception quality",
                "Review tasks clearly identify missing mappings, unsupported formulas, source gaps, or conflicting confirmed facts.",
                "Faster review and fewer surprise issues.",
            ),
            (
                "Control acceptance",
                "SMEs confirm the workflow supports preparer-reviewer control and prevents uncontrolled finalization.",
                "Enterprise adoption readiness.",
            ),
        ],
        [dxa(1.65), dxa(2.45), dxa(2.4)],
    )

    add_heading(doc, "Risk Controls Built Into the Concept", 1)
    for item in [
        "Human-in-the-loop by design: suggested facts, mappings, warnings, and export requests remain reviewable decisions.",
        "Source-grounded reporting: report outputs carry input artifacts, source versions, mapping snapshots, confidence, and evidence.",
        "No black-box finalization: validation readiness and export approval are separate steps.",
        "Data-governance path: pilot can start on synthetic or sanitized files, then move to approved infrastructure only after privacy and security review.",
        "Extensible beyond cash flow: the same repository, mapping, validation, and lineage pattern can support NAV packages, investor reporting, capital account statements, audit support, and regulatory exhibits.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Close", 1)
    add_paragraph(
        doc,
        "The strongest version of this idea is not a standalone AI experiment. It is an Alter Domus operating capability: a governed reporting control tower that combines specialist fund knowledge, repeatable automation, and audit-ready evidence. With a focused pilot, the company can test whether this becomes a reusable internal accelerator for fund teams and a visible proof point for technology-led service quality.",
    )
    add_heading(doc, "Suggested Verbal Pitch", 2)
    add_paragraph(
        doc,
        "We have built the early shape of a fund reporting control tower: a workspace that reads fund documents and datasets, maps them into a governed reporting vocabulary, generates draft client-format workbooks, and shows the evidence behind every number. The important point is the control model. It routes uncertainty to review, records decisions, validates readiness, and blocks final export until approval. I am asking for a small pilot so we can prove whether this can reduce repeat reporting effort while making Alter Domus outputs more consistent, explainable, and audit-ready.",
    )
    add_callout(
        doc,
        "Suggested submission headline",
        "Fund Reporting Control Tower: governed AI assistance for faster, traceable fund administration reporting.",
    )

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    path = build_doc()
    print(path)
